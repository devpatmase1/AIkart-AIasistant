"""
Fast Document Parser (pypdfium2)
================================
Ultra-fast document parser for PDFs, DOCX, TXT, and Markdown files.
Parses multi-page PDFs (e.g. 70+ pages) in under 1 second using pypdfium2 (C++ engine).
"""
from __future__ import annotations

import logging
import time
from pathlib import Path
from typing import Optional

import pypdfium2

from app.services.document_parser.base import BaseDocumentParser
from app.services.chunker import DocumentChunker
from app.services.models.parsed_document import (
    EnrichedChunk,
    ParsedDocument,
)

logger = logging.getLogger(__name__)

_FAST_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".html"}


class FastDocumentParser(BaseDocumentParser):
    """
    High-performance document parser powered by pypdfium2.
    Ideal for large documents where sub-second parsing speed is required.
    """

    parser_name = "fast"

    def __init__(self, workspace_id: int, output_dir: Optional[Path] = None):
        super().__init__(workspace_id, output_dir)
        self.chunker = DocumentChunker(chunk_size=500, chunk_overlap=50)

    @staticmethod
    def supported_extensions() -> set[str]:
        return _FAST_EXTENSIONS

    def parse(
        self,
        file_path: str | Path,
        document_id: int,
        original_filename: str,
    ) -> ParsedDocument:
        path = Path(file_path)
        suffix = path.suffix.lower()
        start_time = time.time()

        if suffix == ".pdf":
            parsed = self._parse_pdf(path, document_id, original_filename)
        elif suffix == ".docx":
            parsed = self._parse_docx(path, document_id, original_filename)
        elif suffix in {".txt", ".md", ".html"}:
            parsed = self._parse_text(path, document_id, original_filename)
        else:
            raise ValueError(f"Unsupported file format for FastParser: {suffix}")

        elapsed_ms = int((time.time() - start_time) * 1000)
        logger.info(
            f"[fast] Parsed document {document_id} ({original_filename}) in {elapsed_ms}ms: "
            f"{parsed.page_count} pages, {len(parsed.chunks)} chunks"
        )
        return parsed

    def _parse_pdf(
        self,
        file_path: Path,
        document_id: int,
        original_filename: str,
    ) -> ParsedDocument:
        """Parse PDF using pypdfium2 (sub-second extraction)."""
        pdf = pypdfium2.PdfDocument(str(file_path))
        page_count = len(pdf)
        
        page_texts: list[tuple[int, str]] = []
        full_markdown_parts: list[str] = []

        for i in range(page_count):
            page = pdf[i]
            textpage = page.get_textpage()
            txt = textpage.get_text_range() or ""
            if txt.strip():
                page_texts.append((i + 1, txt))
                full_markdown_parts.append(f"## Page {i + 1}\n\n{txt}")

        full_markdown = "\n\n".join(full_markdown_parts)

        # Build chunks with page metadata
        chunks: list[EnrichedChunk] = []
        chunk_idx = 0

        for page_no, text in page_texts:
            text_chunks = self.chunker.split_text(text, source=original_filename)
            for c in text_chunks:
                enriched = EnrichedChunk(
                    content=c.content,
                    chunk_index=chunk_idx,
                    source_file=original_filename,
                    document_id=document_id,
                    page_no=page_no,
                    heading_path=[f"Page {page_no}"],
                    image_refs=[],
                    table_refs=[],
                    has_table=False,
                    has_code=False,
                )
                chunks.append(enriched)
                chunk_idx += 1

        return ParsedDocument(
            document_id=document_id,
            original_filename=original_filename,
            markdown=full_markdown,
            page_count=page_count,
            chunks=chunks,
            images=[],
            tables=[],
            tables_count=0,
        )

    def _parse_docx(
        self,
        file_path: Path,
        document_id: int,
        original_filename: str,
    ) -> ParsedDocument:
        """Parse DOCX file."""
        import docx
        doc = docx.Document(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)

        text_chunks = self.chunker.split_text(full_text, source=original_filename)
        chunks = [
            EnrichedChunk(
                content=c.content,
                chunk_index=idx,
                source_file=original_filename,
                document_id=document_id,
                page_no=1,
                heading_path=[],
                image_refs=[],
                table_refs=[],
                has_table=False,
                has_code=False,
            )
            for idx, c in enumerate(text_chunks)
        ]

        return ParsedDocument(
            document_id=document_id,
            original_filename=original_filename,
            markdown=full_text,
            page_count=1,
            chunks=chunks,
            images=[],
            tables=[],
            tables_count=0,
        )

    def _parse_text(
        self,
        file_path: Path,
        document_id: int,
        original_filename: str,
    ) -> ParsedDocument:
        """Parse plain text / MD file."""
        content = file_path.read_text(encoding="utf-8", errors="ignore")
        text_chunks = self.chunker.split_text(content, source=original_filename)

        chunks = [
            EnrichedChunk(
                content=c.content,
                chunk_index=idx,
                source_file=original_filename,
                document_id=document_id,
                page_no=1,
                heading_path=[],
                image_refs=[],
                table_refs=[],
                has_table=False,
                has_code=False,
            )
            for idx, c in enumerate(text_chunks)
        ]

        return ParsedDocument(
            document_id=document_id,
            original_filename=original_filename,
            markdown=content,
            page_count=1,
            chunks=chunks,
            images=[],
            tables=[],
            tables_count=0,
        )
